from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

from docx import Document

from backend.agents.graph import build_part2_graph
from backend.agents.nodes_part2 import parse_jd_node
from backend.agents.schemas import JDStruct
from backend.agents.schemas import GraphState


def _create_sample_resume(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane@example.com | +1 555-0101")
    doc.add_paragraph("Projects")
    doc.add_paragraph("ResumeFit")
    doc.add_paragraph("- Built FastAPI endpoints for analysis workflow")
    doc.add_paragraph("Skills")
    doc.add_paragraph("Python, FastAPI, SQL")
    doc.add_paragraph("Education")
    doc.add_paragraph("B.S. Computer Science")
    doc.save(path)


def test_parse_resume_and_jd_into_shared_state(tmp_path: Path) -> None:
    resume_path = tmp_path / "resume.docx"
    _create_sample_resume(resume_path)

    state = GraphState(
        resume_path=str(resume_path),
        jd_text=(
            "Senior Backend Engineer at Acme\n"
            "Requirements: Python, FastAPI, Docker\n"
            "Nice to have: Kubernetes\n"
            "- Build APIs\n"
        ),
    )

    graph = build_part2_graph()
    result = graph.invoke(state)

    assert not [error for error in result["errors"] if "failed" in error]
    assert result["resume_struct"] is not None
    assert result["jd_struct"] is not None

    resume_struct = result["resume_struct"]
    jd_struct = result["jd_struct"]

    assert len(resume_struct.sections) >= 3
    assert any(section.title.lower() == "projects" for section in resume_struct.sections)
    assert "Python" in resume_struct.skills

    assert "python" in [item.lower() for item in jd_struct.required_skills]
    assert any("kubernetes" in item.lower() for item in jd_struct.nice_to_haves)


def test_corrupt_resume_is_handled_gracefully(tmp_path: Path) -> None:
    corrupt_path = tmp_path / "bad.docx"
    corrupt_path.write_text("not a real docx", encoding="utf-8")

    state = GraphState(resume_path=str(corrupt_path), jd_text="Backend Engineer role")
    graph = build_part2_graph()
    result = graph.invoke(state)

    assert result.get("resume_struct") is None
    assert any("parse_resume failed" in err for err in result["errors"])


def test_parse_jd_node_adds_warning_when_fallback_used() -> None:
    state = GraphState(resume_path="", jd_text="Requirements: Python")
    fake_struct = JDStruct(required_skills=["Python"])

    with patch("backend.services.jd_parser.parse_jd", return_value=(fake_struct, "fallback")):
        updated = parse_jd_node(state)

    assert updated.jd_struct is not None
    assert any("fallback parser" in err for err in updated.errors)
