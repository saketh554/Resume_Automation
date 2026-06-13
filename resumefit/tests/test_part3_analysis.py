from __future__ import annotations

from pathlib import Path

from docx import Document

from backend.agents.graph import build_agent_graph
from backend.agents.schemas import GraphState


def _create_resume_for_part3(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane@example.com")
    doc.add_paragraph("Projects")
    doc.add_paragraph("ResumeFit")
    doc.add_paragraph("- Built FastAPI endpoints and SQL data layer")
    doc.add_paragraph("- Responsible for various tasks in backend operations")
    doc.add_paragraph("Skills")
    doc.add_paragraph("Python, FastAPI, SQL")
    doc.add_paragraph("Education")
    doc.add_paragraph("B.S. Computer Science")
    doc.save(path)


def test_part3_alignment_formula_is_reproducible(tmp_path: Path) -> None:
    resume_path = tmp_path / "resume.docx"
    _create_resume_for_part3(resume_path)

    state = GraphState(
        resume_path=str(resume_path),
        jd_text=(
            "Senior Backend Engineer\n"
            "Requirements: Python, FastAPI, Docker\n"
            "Nice to have: Kubernetes\n"
        ),
    )

    graph = build_agent_graph()
    result_a = graph.invoke(state)
    result_b = graph.invoke(state)

    assert result_a["alignment"] is not None
    assert result_b["alignment"] is not None
    assert result_a["alignment"].alignment_pct == result_b["alignment"].alignment_pct

    # Required matched: python + fastapi = 2/3; nice matched: 0/1.
    # Score = (2/3)*0.7 + 0*0.3 = 0.466666... => 46.67
    assert result_a["alignment"].alignment_pct == 46.67


def test_missing_skill_has_one_bullet_and_project_mapping(tmp_path: Path) -> None:
    resume_path = tmp_path / "resume.docx"
    _create_resume_for_part3(resume_path)

    state = GraphState(
        resume_path=str(resume_path),
        jd_text="Requirements: Python, FastAPI, Docker\nNice to have: Kubernetes",
    )
    graph = build_agent_graph()
    result = graph.invoke(state)

    alignment = result["alignment"]
    gaps = result["gap_bullets"]
    mappings = result["project_mappings"]

    assert alignment is not None
    assert len(gaps) == len(alignment.missing_skills)
    assert len(gaps) > 0

    mapping_by_skill = {item.skill: item for item in mappings}
    for gap in gaps:
        assert gap.skill in mapping_by_skill
        assert len(mapping_by_skill[gap.skill].project_names) >= 1
        assert "[X]%" in gap.bullet


def test_audit_only_flags_project_experience_content(tmp_path: Path) -> None:
    resume_path = tmp_path / "resume.docx"
    _create_resume_for_part3(resume_path)

    state = GraphState(
        resume_path=str(resume_path),
        jd_text="Requirements: Python",
    )
    graph = build_agent_graph()
    result = graph.invoke(state)

    audit = result["project_audit"]
    assert len(audit) >= 1
    assert all(item.project_name.lower() in {"projects", "experience", "work experience"} for item in audit)
    assert not any(item.project_name.lower() == "contact" for item in audit)

    combined_flags = [point.lower() for item in audit for point in item.irrelevant_points]
    assert any("various tasks" in point for point in combined_flags)
