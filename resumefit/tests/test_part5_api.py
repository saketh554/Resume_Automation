from __future__ import annotations

import importlib
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient


def _create_resume(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("JANE DOE")
    doc.add_paragraph("jane@example.com")
    doc.add_paragraph("PROJECTS")
    doc.add_paragraph("ResumeFit")
    doc.add_paragraph("- Built FastAPI endpoints and SQL data layer")
    doc.add_paragraph("- Responsible for various tasks in backend operations")
    doc.add_paragraph("SKILLS")
    doc.add_paragraph("Python, FastAPI, SQL")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("B.S. Computer Science")
    doc.save(path)


def test_part5_api_flow_end_to_end(tmp_path: Path) -> None:
    import backend.app.main as app_main

    importlib.reload(app_main)

    resume_path = tmp_path / "resume.docx"
    _create_resume(resume_path)

    with TestClient(app_main.app) as client:
        with resume_path.open("rb") as f:
            upload = client.post(
                "/api/resume",
                files={"file": ("resume.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
        assert upload.status_code == 200
        resume_id = upload.json()["resume_id"]

        analyze = client.post(
            "/api/analyze",
            json={
                "resume_id": resume_id,
                "jd_text": "Senior Backend Engineer at Acme\nRequirements: Python, FastAPI, Docker\nNice to have: Kubernetes",
            },
        )
        assert analyze.status_code == 200
        analysis_id = analyze.json()["analysis_id"]

        detail = client.get(f"/api/analysis/{analysis_id}")
        assert detail.status_code == 200
        assert detail.json()["analysis_id"] == analysis_id

        resume_download = client.get(f"/api/analysis/{analysis_id}/resume.docx")
        assert resume_download.status_code == 200
        assert resume_download.headers["content-type"].startswith(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        email = client.get(f"/api/analysis/{analysis_id}/email")
        assert email.status_code == 200
        email_json = email.json()
        assert "subject" in email_json and email_json["subject"]
        assert "body" in email_json and email_json["body"]
