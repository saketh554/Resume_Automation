from __future__ import annotations

from pathlib import Path

from docx import Document

from backend.agents.graph import build_agent_graph
from backend.agents.schemas import GraphState


def _create_resume_for_part4(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("JANE DOE")
    doc.add_paragraph("jane@example.com | +1 555-0101")
    doc.add_paragraph("PROJECTS")
    doc.add_paragraph("ResumeFit")
    doc.add_paragraph("- Built FastAPI endpoints and SQL data layer")
    doc.add_paragraph("- Responsible for various tasks in backend operations")
    doc.add_paragraph("SKILLS")
    doc.add_paragraph("Python, FastAPI, SQL")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("B.S. Computer Science")
    doc.save(path)


def test_part4_generates_tailored_docx_with_structure_preserved(tmp_path: Path) -> None:
    resume_path = tmp_path / "resume.docx"
    _create_resume_for_part4(resume_path)

    graph = build_agent_graph()
    state = GraphState(
        resume_path=str(resume_path),
        jd_text=(
            "Senior Backend Engineer at Acme\n"
            "Requirements: Python, FastAPI, Docker\n"
            "Nice to have: Kubernetes\n"
        ),
    )

    result = graph.invoke(state)
    assert result["tailored_docx_path"]

    original = Document(str(resume_path))
    tailored = Document(result["tailored_docx_path"])

    original_headers = [p.text.strip() for p in original.paragraphs if p.text.strip().isupper()]
    tailored_headers = [p.text.strip() for p in tailored.paragraphs if p.text.strip().isupper()]

    assert original_headers == tailored_headers
    assert abs(len(original.paragraphs) - len(tailored.paragraphs)) <= 1

    tailored_text = "\n".join(p.text for p in tailored.paragraphs)
    assert "[X]%" in tailored_text


def test_part4_cold_email_is_grounded_and_non_fabricated(tmp_path: Path) -> None:
    resume_path = tmp_path / "resume.docx"
    _create_resume_for_part4(resume_path)

    graph = build_agent_graph()
    state = GraphState(
        resume_path=str(resume_path),
        jd_text=(
            "Senior Backend Engineer at Acme\n"
            "Requirements: Python, FastAPI\n"
            "Nice to have: Docker\n"
        ),
    )

    result = graph.invoke(state)
    email = result.get("cold_email")
    assert email is not None

    combined = f"{email.subject}\n{email.body}".lower()
    assert "acme" in combined or "your team" in combined
    assert "senior backend engineer" in combined or "the role" in combined

    # Fabrication guardrails for known forbidden fields in this fixture.
    assert "phd" not in combined
    assert "stanford" not in combined
