from __future__ import annotations

from pathlib import Path

from docx import Document

from backend.agents.schemas import ProjectAuditItem, ProjectMapping
from backend.app.settings import settings


def _copy_paragraph_format(source_para, target_para) -> None:
    target_para.style = source_para.style
    target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
    target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
    target_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
    target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
    target_para.paragraph_format.space_after = source_para.paragraph_format.space_after
    target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing


def _copy_run_style(source_run, target_run) -> None:
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size


def _section_titles(paragraph_texts: list[str]) -> list[str]:
    return [line.strip().lower().rstrip(":") for line in paragraph_texts if line.strip().isupper() or line.strip().endswith(":")]


def apply_docx_edits(
    source_path: str,
    mappings: list[ProjectMapping],
    audits: list[ProjectAuditItem],
) -> str:
    source = Path(source_path)
    if not source.exists():
        raise ValueError("Source resume .docx does not exist.")

    output_dir = Path(settings.data_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document(source)
    original_paragraph_count = len(doc.paragraphs)
    original_titles = _section_titles([p.text for p in doc.paragraphs])
    original_styles = {style.name for style in doc.styles}

    bullets_to_add = [item.bullet for item in mappings]
    removable_points = [point for audit in audits for point in audit.irrelevant_points]

    # One-in-one-out: first replace audited points in-place to preserve paragraph count.
    replaced_bullets: list[str] = []
    bullet_cursor = 0
    for point in removable_points:
        if bullet_cursor >= len(bullets_to_add):
            break
        for para in doc.paragraphs:
            if point and point in para.text:
                para.text = f"- {bullets_to_add[bullet_cursor]}"
                replaced_bullets.append(bullets_to_add[bullet_cursor])
                bullet_cursor += 1
                break

    remaining_bullets = bullets_to_add[bullet_cursor:]

    # Insert new bullets by cloning style from an existing bullet-like paragraph.
    template_para = None
    for para in doc.paragraphs:
        if para.text.strip().startswith(('-', '*')):
            template_para = para
            break
    if template_para is None:
        template_para = doc.paragraphs[-1]

    for bullet in remaining_bullets:
        new_para = doc.add_paragraph("")
        _copy_paragraph_format(template_para, new_para)
        source_run = template_para.runs[0] if template_para.runs else None
        new_run = new_para.add_run(f"- {bullet}")
        if source_run is not None:
            _copy_run_style(source_run, new_run)

    output_path = output_dir / f"tailored_{source.stem}.docx"
    doc.save(output_path)

    # Post-edit assertions
    check_doc = Document(output_path)
    updated_paragraph_count = len(check_doc.paragraphs)
    updated_titles = _section_titles([p.text for p in check_doc.paragraphs])
    updated_styles = {style.name for style in check_doc.styles}

    if original_titles != updated_titles:
        raise ValueError("Layout drift detected: section heading order changed.")

    if not original_styles.issubset(updated_styles):
        raise ValueError("Layout drift detected: styles changed after edit.")

    if abs(updated_paragraph_count - original_paragraph_count) > 1:
        raise ValueError("Layout drift detected: paragraph count outside tolerance.")

    return str(output_path)
